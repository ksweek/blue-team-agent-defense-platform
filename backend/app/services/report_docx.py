from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"

TONE_COLORS: dict[str, dict[str, str]] = {
    "safe": {"text": "0B6B46", "fill": "DFF4EA", "accent": "1F9D68"},
    "warn": {"text": "9A5D00", "fill": "FFF1D6", "accent": "D08A00"},
    "danger": {"text": "A12F27", "fill": "FFE2DD", "accent": "D65745"},
    "info": {"text": "1D5FBF", "fill": "E4EEFF", "accent": "4A7FE2"},
}


def build_report_docx_artifact(payload: dict[str, Any], output_path: Path) -> None:
    document = Document()
    _configure_document(document)
    _set_core_properties(document, payload)
    _configure_header_footer(document, payload)

    _add_cover_page(document, payload)
    _add_summary_section(document, payload)
    _add_object_section(document, payload)
    _add_trace_section(document, payload)
    _add_payload_section(document, payload)
    _add_sensitive_section(document, payload)
    _add_timeline_section(document, payload)
    _add_recommendation_section(document, payload)
    _add_appendix_section(document, payload)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _configure_document(document: DocumentType) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(0.9)
    section.start_type = WD_SECTION.NEW_PAGE

    for style_name, size in (
        ("Normal", 10.5),
        ("Title", 22),
        ("Heading 1", 15),
        ("Heading 2", 12.5),
        ("Heading 3", 11),
    ):
        style = document.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        _set_style_east_asia_font(style, BODY_FONT)


def _set_core_properties(document: DocumentType, payload: dict[str, Any]) -> None:
    report = payload.get("report") or {}
    presentation = payload.get("presentation") or {}
    properties = document.core_properties
    properties.title = str(report.get("report_name") or "中文安全报告")
    properties.subject = str(presentation.get("report_type_label") or "安全报告")
    properties.author = "蓝队防御平台"
    properties.comments = str(presentation.get("summary_text") or "")
    properties.category = "Security Report"


def _configure_header_footer(document: DocumentType, payload: dict[str, Any]) -> None:
    report = payload.get("report") or {}
    task = payload.get("task") or {}
    for section in document.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _clear_paragraph(header_para)
        _add_run(
            header_para,
            f"蓝队防御平台 | {report.get('report_name') or '安全报告'} | 任务 #{task.get('id') or '-'}",
            size=8.5,
            color="5D718A",
        )

        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_paragraph(footer_para)
        _add_run(footer_para, "内部安全评估材料  第 ", size=8.5, color="6A7B90")
        _append_field(footer_para, "PAGE")
        _add_run(footer_para, " 页", size=8.5, color="6A7B90")


def _add_cover_page(document: DocumentType, payload: dict[str, Any]) -> None:
    report = payload.get("report") or {}
    task = payload.get("task") or {}
    event = payload.get("event") or {}
    template = payload.get("template") or {}
    presentation = payload.get("presentation") or {}

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(kicker, "蓝队防御平台 / 中文正式报告", size=10.5, bold=True, color="2D5FA8")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(6)
    _add_run(title, "AI 安全防护分析报告", size=24, bold=True, color="12263F")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    _add_run(
        subtitle,
        f"{presentation.get('report_type_label') or '安全报告'} | 报告编号 #{report.get('id') or '-'}",
        size=11,
        color="5D718A",
    )

    raw_name = document.add_paragraph()
    raw_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    raw_name.paragraph_format.space_after = Pt(14)
    _add_run(raw_name, f"内部报告名：{report.get('report_name') or '-'}", size=10, color="7A8CA3")

    summary_box = document.add_table(rows=1, cols=1)
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_box.autofit = True
    summary_cell = summary_box.cell(0, 0)
    _set_cell_shading(summary_cell, "EEF4FF")
    _set_cell_margins(summary_cell, 180, 120, 180, 120)
    _write_cell_title(summary_cell, "执行摘要", color="2D5FA8")
    _append_cell_paragraph(summary_cell, str(presentation.get("summary_text") or "-"), size=10.5, color="12263F")
    _append_cell_paragraph(
        summary_cell,
        "详细原始返回、命中证据和敏感数据痕迹见正文与附录。",
        size=9.5,
        color="5D718A",
    )

    document.add_paragraph()
    _add_highlight_grid(document, presentation.get("highlights") or [])
    document.add_paragraph()
    _add_metadata_table(
        document,
        [
            ("模板版本", str(template.get("template_version") or "-")),
            ("模板名称", str(template.get("template_name") or "-")),
            ("报告类型", str(presentation.get("report_type_label") or "-")),
            ("内部报告名", str(report.get("report_name") or "-")),
            ("任务名称", str(task.get("task_name") or "-")),
            ("目标对象", str(task.get("target_agent") or "-")),
            ("事件编号", str(event.get("id") or "-")),
            ("报告生成时间", str(report.get("created_at") or "-")),
            ("报告导出时间", str(report.get("exported_at") or "-")),
        ],
    )
    document.add_page_break()


def _add_summary_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    decision_summary = presentation.get("decision_summary") or {}
    _section_heading(document, "一、执行摘要")
    _callout_block(
        document,
        str(decision_summary.get("title") or "关键结论"),
        str(decision_summary.get("summary") or "-"),
        str(decision_summary.get("detail") or ""),
        "info",
    )

    trace_items = presentation.get("trace_summary") or []
    if trace_items:
        _sub_heading(document, "1.1 判定摘要")
        table = _create_key_value_table(document, 3, ("项目", "结论", "说明"))
        for item in trace_items:
            row = table.add_row().cells
            _fill_table_row(
                row,
                (
                    str(item.get("label") or "-"),
                    str(item.get("value") or "-"),
                    str(item.get("detail") or "-"),
                ),
                tone=str(item.get("tone") or "info"),
            )


def _add_object_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    _section_heading(document, "二、对象与任务信息")
    items = presentation.get("object_summary") or []
    _add_metadata_table(document, [(str(key), str(value)) for key, value in items])


def _add_trace_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    _section_heading(document, "三、执行链判断")

    _sub_heading(document, "3.1 授权与规则链路")
    trace_items = presentation.get("trace_summary") or []
    if trace_items:
        for item in trace_items:
            _callout_block(
                document,
                str(item.get("label") or "-"),
                str(item.get("value") or "-"),
                str(item.get("detail") or ""),
                str(item.get("tone") or "info"),
                compact=True,
            )
    else:
        _add_body_paragraph(document, "当前没有记录到执行链判断信息。", color="5D718A")

    _sub_heading(document, "3.2 命中概览")
    _add_bullet_group(document, "命中控制面", presentation.get("matched_controls") or [])
    _add_bullet_group(document, "命中规则", presentation.get("matched_rules") or [])
    _add_bullet_group(document, "攻击信号", presentation.get("matched_signals") or [])


def _add_payload_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    items = presentation.get("payload_hits") or []
    _section_heading(document, "四、重点命中与证据")

    if not items:
        _add_body_paragraph(document, "当前没有可展示的关键命中项。", color="5D718A")
        return

    for index, item in enumerate(items, start=1):
        _sub_heading(document, f"4.{index} {item.get('display_label') or item.get('label') or '-'}")
        table = _create_key_value_table(document, 2, ("字段", "内容"))
        _fill_table_row(table.add_row().cells, ("命中类型", str(item.get("kind") or "-")), tone=_payload_tone(str(item.get("kind") or "")))
        _fill_table_row(table.add_row().cells, ("来源", str(item.get("source_label") or item.get("source") or "-")))
        _fill_table_row(table.add_row().cells, ("位置", str(item.get("location_label") or "正文")))
        if item.get("category_label"):
            _fill_table_row(table.add_row().cells, ("规则分类", str(item.get("category_label") or "-")))
        _fill_table_row(table.add_row().cells, ("说明", str(item.get("detail_label") or item.get("detail") or "-")))
        evidence = str(item.get("evidence") or "").strip()
        if evidence:
            _add_code_block(document, evidence, title="证据片段")


def _add_sensitive_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    items = presentation.get("sensitive_hits") or []
    _section_heading(document, "五、敏感数据痕迹")

    if not items:
        _add_body_paragraph(document, "当前没有识别到敏感数据痕迹。", color="5D718A")
        return

    table = _create_key_value_table(document, 4, ("分类", "来源", "位置", "脱敏预览"))
    for item in items:
        _fill_table_row(
            table.add_row().cells,
            (
                str(item.get("label") or "-"),
                str(item.get("source") or "-"),
                str(item.get("location_label") or "正文"),
                str(item.get("preview") or "-"),
            ),
            tone="warn",
        )


def _add_timeline_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    items = presentation.get("timeline") or []
    _section_heading(document, "六、执行时间线")

    if not items:
        _add_body_paragraph(document, "当前没有时间线数据。", color="5D718A")
        return

    table = _create_key_value_table(document, 3, ("节点", "时间", "说明"))
    for item in items:
        _fill_table_row(
            table.add_row().cells,
            (
                str(item.get("label") or "-"),
                str(item.get("value") or "-"),
                str(item.get("detail") or "-"),
            ),
            tone="info",
        )


def _add_recommendation_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    items = presentation.get("recommendations") or []
    _section_heading(document, "七、处置建议")

    if not items:
        _add_body_paragraph(document, "当前没有处置建议。", color="5D718A")
        return

    for index, item in enumerate(items, start=1):
        tone = str(item.get("tone") or "info")
        _callout_block(
            document,
            f"建议 {index}：{item.get('title') or '-'}",
            str(item.get("detail") or "-"),
            "",
            tone,
            compact=True,
        )


def _add_appendix_section(document: DocumentType, payload: dict[str, Any]) -> None:
    presentation = payload.get("presentation") or {}
    items = presentation.get("appendix_sections") or []
    _section_heading(document, "八、原始数据附录")

    if not items:
        _add_body_paragraph(document, "当前没有附录数据。", color="5D718A")
        return

    for index, item in enumerate(items, start=1):
        _sub_heading(document, f"8.{index} {item.get('title') or '-'}")
        _add_code_block(document, _serialize_block(item.get("content")), mono=True)


def _add_highlight_grid(document: DocumentType, items: list[dict[str, Any]]) -> None:
    if not items:
        return

    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for index, item in enumerate(items[:4]):
        cell = table.cell(index // 2, index % 2)
        tone = str(item.get("tone") or "info")
        colors = TONE_COLORS.get(tone, TONE_COLORS["info"])
        _set_cell_shading(cell, colors["fill"])
        _set_cell_margins(cell, 180, 120, 180, 120)
        _write_cell_title(cell, str(item.get("label") or "-"), color=colors["accent"])
        _append_cell_paragraph(cell, str(item.get("value") or "-"), size=14, bold=True, color=colors["text"])
        _append_cell_paragraph(cell, str(item.get("detail") or ""), size=9, color="5D718A")


def _add_metadata_table(document: DocumentType, rows: list[tuple[str, str]]) -> None:
    table = _create_key_value_table(document, 2, ("字段", "内容"))
    for key, value in rows:
        _fill_table_row(table.add_row().cells, (key, value))


def _add_bullet_group(document: DocumentType, title: str, items: list[str]) -> None:
    label = document.add_paragraph()
    label.paragraph_format.space_before = Pt(4)
    label.paragraph_format.space_after = Pt(1)
    _add_run(label, title, bold=True, size=10.5, color="12263F")

    if not items:
        _add_body_paragraph(document, "未记录。", color="5D718A")
        return

    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(0.5)
        _add_run(paragraph, str(item), size=10.2, color="12263F")


def _callout_block(
    document: DocumentType,
    title: str,
    summary: str,
    detail: str,
    tone: str,
    *,
    compact: bool = False,
) -> None:
    colors = TONE_COLORS.get(tone, TONE_COLORS["info"])
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, colors["fill"])
    _set_cell_margins(cell, 180, 120, 180, 120)
    _write_cell_title(cell, title, color=colors["accent"])
    _append_cell_paragraph(cell, summary, size=11 if not compact else 10.5, bold=not compact, color=colors["text"])
    if detail:
        _append_cell_paragraph(cell, detail, size=9.6, color="5D718A")


def _add_code_block(document: DocumentType, text: str, *, title: str | None = None, mono: bool = True) -> None:
    if title:
        title_para = document.add_paragraph()
        title_para.paragraph_format.space_before = Pt(4)
        title_para.paragraph_format.space_after = Pt(1)
        _add_run(title_para, title, bold=True, size=10, color="5D718A")

    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "F5F7FB")
    _set_cell_margins(cell, 160, 90, 160, 90)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = _add_run(paragraph, text, size=8.8, color="2E4057", font_name=MONO_FONT if mono else BODY_FONT)
    run.add_break(WD_BREAK.LINE)


def _section_heading(document: DocumentType, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(5)
    _add_run(paragraph, text, size=15, bold=True, color="12263F")


def _sub_heading(document: DocumentType, text: str) -> None:
    paragraph = document.add_paragraph(style="Heading 2")
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)
    _add_run(paragraph, text, size=11.5, bold=True, color="223A56")


def _add_body_paragraph(document: DocumentType, text: str, *, color: str = "12263F") -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.3
    _add_run(paragraph, text, size=10.5, color=color)


def _create_key_value_table(document: DocumentType, cols: int, headers: tuple[str, ...]):
    table = document.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header_cells = table.rows[0].cells
    for cell, header in zip(header_cells, headers):
        _set_cell_shading(cell, "EAF1FB")
        _set_cell_margins(cell, 120, 90, 120, 90)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(paragraph, header, size=9.6, bold=True, color="2A466A")
    return table


def _fill_table_row(cells: Any, values: tuple[str, ...], tone: str | None = None) -> None:
    for index, (cell, value) in enumerate(zip(cells, values)):
        _set_cell_margins(cell, 120, 90, 120, 90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        if tone and index == 0 and len(values) > 2:
            colors = TONE_COLORS.get(tone, TONE_COLORS["info"])
            _set_cell_shading(cell, colors["fill"])
        paragraph = cell.paragraphs[0]
        _clear_paragraph(paragraph)
        _add_run(paragraph, value, size=9.8, color="12263F")


def _payload_tone(kind: str) -> str:
    if kind == "signal":
        return "danger"
    if kind == "rule":
        return "warn"
    if kind == "control":
        return "safe"
    return "info"


def _serialize_block(value: Any) -> str:
    if value in (None, "", [], {}):
        return "没有可展示的数据。"
    if isinstance(value, str):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError:
            return value
        return json.dumps(parsed, ensure_ascii=False, indent=2)
    return json.dumps(value, ensure_ascii=False, indent=2)


def _write_cell_title(cell: Any, text: str, *, color: str) -> None:
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.paragraph_format.space_after = Pt(1)
    _add_run(paragraph, text, size=9.5, bold=True, color=color)


def _append_cell_paragraph(
    cell: Any,
    text: str,
    *,
    size: float,
    color: str,
    bold: bool = False,
) -> None:
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0.5)
    paragraph.paragraph_format.line_spacing = 1.2
    _add_run(paragraph, text, size=size, bold=bold, color=color)


def _add_run(
    paragraph: Any,
    text: str,
    *,
    size: float,
    color: str,
    bold: bool = False,
    font_name: str = BODY_FONT,
):
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.color.rgb = RGBColor.from_string(color)
    _set_run_east_asia_font(run, font_name)
    return run


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell: Any, left: int, top: int, right: int, bottom: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("left", left), ("top", top), ("right", right), ("bottom", bottom)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_style_east_asia_font(style: Any, font_name: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def _set_run_east_asia_font(run: Any, font_name: str) -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def _clear_paragraph(paragraph: Any) -> None:
    for element in list(paragraph._element):
        paragraph._element.remove(element)


def _append_field(paragraph: Any, field_name: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    paragraph._p.append(begin)
    paragraph._p.append(instr)
    paragraph._p.append(separate)
    paragraph._p.append(end)
